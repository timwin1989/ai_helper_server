from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, ListFlowable, ListItem, Paragraph, Spacer, Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from pathlib import Path
from bs4 import BeautifulSoup
# from xhtml2pdf import pisa
import datetime

def save_html_to_pdf(html_string: str, filename: str = None):
    font_path = "fonts/DejaVuSans.ttf"
    header_path = "assets/header_pdf.png"

    if not Path(header_path).exists():
        print(f"❌ Картинка не найдена: {header_path}")
        return
    if not Path(font_path).exists():
        print(f"❌ Шрифт не найден: {font_path}")
        return

    pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))

    Path("answers").mkdir(parents=True, exist_ok=True)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = f"answers/{filename or f'report_{timestamp}.pdf'}"

    doc = SimpleDocTemplate(output_name, pagesize=A4,
                            leftMargin=50, rightMargin=40, topMargin=30, bottomMargin=50)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='KazakhText', fontName='DejaVuSans', fontSize=11, leading=16))

    elements = []

    # Картинка
    width, _ = A4
    img = Image(header_path, width=width - 90, height=150)
    elements.append(img)
    elements.append(Spacer(1, 12))

    # Обработка HTML
    cleaned_text = html_string.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    soup = BeautifulSoup(cleaned_text, "html.parser")

    for tag in soup.children:
        if tag.name == "p":
            text = tag.get_text(strip=True)
            if text:
                elements.append(Paragraph(text, styles['KazakhText']))
                elements.append(Spacer(1, 8))
        elif tag.name in ["ol", "ul"]:
            items = []
            for li in tag.find_all("li"):
                li_text = li.get_text(strip=True)
                if li_text:
                    items.append(ListItem(Paragraph(li_text, styles['KazakhText'])))
            if items:
                elements.append(ListFlowable(items, bulletType='1' if tag.name == "ol" else 'bullet'))
                elements.append(Spacer(1, 8))

    doc.build(elements)
    print(f"✅ PDF сохранён: {output_name}")


# def save_to_pdf(text: str, filename: str = None):
#     font_path = "fonts/DejaVuSans.ttf"
#     header_path = "assets/header_pdf.png"  # ⚠️ путь к картинке
#     if not Path(font_path).exists():
#         print(f"❌ Шрифт не найден: {font_path}")
#         return
#     if not Path(header_path).exists():
#         print(f"❌ Картинка не найдена: {header_path}")
#         return

#     pdfmetrics.registerFont(TTFont("DejaVu", font_path))

#     timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
#     #Path("answers").mkdir(parents=True, exist_ok=True)  # создаёт папку, если нет
#     output_name = f"answers/{filename or f'generated_answer_{timestamp}.pdf'}"

#     doc = SimpleDocTemplate(output_name, pagesize=A4,
#                             leftMargin=50, rightMargin=40, topMargin=50, bottomMargin=50)

#     styles = getSampleStyleSheet()
#     styles.add(ParagraphStyle(name='KazakhText', fontName='DejaVu', fontSize=11, leading=16))

#     elements = []

#     # === Добавляем картинку шапки ===
#     width, height = A4
#     img = Image(header_path, width=width - 90, height=80)  # чуть меньше ширины страницы
#     elements.append(img)
#     elements.append(Spacer(1, 12))

#     # === Основной текст ===
#     for paragraph in text.split("\n"):
#         paragraph = paragraph.strip()
#         if not paragraph:
#             elements.append(Spacer(1, 12))
#             continue

#         # стиль с отступом для подабзацев
#         if paragraph.lstrip().startswith(("a)", "b)", "c)", "d)", "-", "•", "*")):
#             style = ParagraphStyle('ListItem', parent=styles['KazakhText'], leftIndent=20)
#         else:
#             style = styles['KazakhText']

#         elements.append(Paragraph(paragraph, style))

#     doc.build(elements)
#     print(f"\n📄 PDF успешно создан с шапкой: {output_name}")

def save_to_docx(text: str, filename: str = None):
    header_path = "assets/header_docx.png"  # отдельное изображение (можно то же)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    Path("answers").mkdir(parents=True, exist_ok=True)
    output_name = f"answers/{filename or f'generated_answer_{timestamp}.docx'}"

    doc = Document()

    # === Добавляем шапку-картинку ===
    if Path(header_path).exists():
        doc.add_picture(header_path, width=Inches(6))  # ширина на всю страницу
        doc.add_paragraph()  # отступ

    # === Основной текст ===
    for paragraph in text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
        else:
            doc.add_paragraph("")

    doc.save(output_name)
    print(f"📝 DOCX успешно сохранён: {output_name}")